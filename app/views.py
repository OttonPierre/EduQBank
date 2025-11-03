from django.shortcuts import render, redirect
from .models import Conteudo
from .forms import QuestaoForm
from django.http import JsonResponse, HttpResponse
from django.contrib import messages
from django.contrib.auth.models import User
from django.contrib.auth import authenticate
from rest_framework.decorators import api_view, permission_classes
from rest_framework.response import Response
from rest_framework import status
from rest_framework.permissions import AllowAny
from rest_framework_simplejwt.tokens import RefreshToken
from rest_framework import viewsets
from rest_framework import generics
from .models import Questao
from .serializers import QuestaoSerializer, ConteudoSerializer  
import json
from django.views.decorators.csrf import csrf_exempt
from django.core.files.storage import default_storage
from bs4 import BeautifulSoup
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image as ReportLabImage, KeepTogether
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.platypus.tableofcontents import TableOfContents
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import io
import os
import urllib.parse
from PIL import Image as PILImage
from matplotlib import mathtext
from matplotlib import font_manager
import base64
from .utils import render_latex_to_png_bytes, split_text_and_math, html_render_math_to_img

## moved to utils.py

def _extract_text_from_html(html: str) -> str:
    if not html:
        return ""
    soup = BeautifulSoup(html, 'lxml')
    # Replace <br> with newlines
    for br in soup.find_all("br"):
        br.replace_with("\n")
    text = soup.get_text("\n")
    return text.strip()

def _docx_add_if_header(document: Document, title: str = "PROVA", subtitle: str = "") -> None:
    section = document.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    run = header_para.add_run("INSTITUTO FEDERAL – Sistema de Avaliação\n")
    run.font.size = Pt(10)
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph()
    r = p.add_run(title)
    r.font.size = Pt(16)
    r.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        p2 = document.add_paragraph(subtitle)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info = document.add_paragraph("Professor(a): __________________    Turma: ______    Data: ____/____/______    Nota: ______")
    info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph("Aluno(a): _________________________________________________________________")
    document.add_paragraph("")

def _pdf_if_header(canvas, doc_title: str = "PROVA", subtitle: str = ""):
    width, height = A4
    canvas.setFont("Helvetica-Bold", 10)
    canvas.drawCentredString(width/2, height - 40, "INSTITUTO FEDERAL – Sistema de Avaliação")
    canvas.setFont("Helvetica-Bold", 16)
    canvas.drawCentredString(width/2, height - 70, doc_title)
    if subtitle:
        canvas.setFont("Helvetica", 12)
        canvas.drawCentredString(width/2, height - 88, subtitle)
    canvas.setFont("Helvetica", 10)
    canvas.drawString(40, height - 120, "Professor(a): __________________    Turma: ______    Data: ____/____/______    Nota: ______")
    canvas.drawString(40, height - 140, "Aluno(a): _________________________________________________________________")

def _build_exam_html(questions):
    # Minimal HTML with IF header and MathJax
    head = """
    <head>
      <meta charset='utf-8'/>
      <style>
        body { font-family: Arial, sans-serif; margin: 40px; }
        h1, h2 { text-align: center; margin: 0; }
        .meta { margin-top: 16px; font-size: 12px; }
        .question { margin-top: 18px; }
        .answer { margin-top: 8px; color: #333; }
        img { max-width: 100%; }
      </style>
      <script>
        window.MathJax = { tex: { inlineMath: [['$', '$'], ['\\\(', '\\\)']], displayMath: [['$$','$$'], ['\\\[','\\\]']] } };
      </script>
      <script src='https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-svg.js'></script>
    </head>
    """
    body_parts = [
        "<h2>INSTITUTO FEDERAL – Sistema de Avaliação</h2>",
        "<h1>PROVA</h1>",
        "<div class='meta'>Professor(a): __________________ &nbsp;&nbsp; Turma: ______ &nbsp;&nbsp; Data: ____/____/______ &nbsp;&nbsp; Nota: ______</div>",
        "<div class='meta'>Aluno(a): _________________________________________________________________</div>",
    ]
    for idx, q in enumerate(questions, start=1):
        body_parts.append(f"<div class='question'><strong>Questão {idx}</strong></div>")
        body_parts.append(f"<div class='question-content'>{q.enunciado or ''}</div>")
        if getattr(q, 'resposta', None):
            body_parts.append("<div class='answer'><em>Resposta:</em></div>")
            body_parts.append(f"<div class='answer-content'>{q.resposta or ''}</div>")
    html = f"<html>{head}<body>{''.join(body_parts)}</body></html>"
    return html

    # Legacy helper retained (no longer used)

def _build_questions_flow(questions):
    styles = getSampleStyleSheet()
    normal = ParagraphStyle(name='Normal', parent=styles['Normal'], fontName='Helvetica', fontSize=11, leading=14, alignment=TA_JUSTIFY)
    title_style = ParagraphStyle(name='QTitle', parent=styles['Heading4'], spaceAfter=6)
    flow = []
    for idx, q in enumerate(questions, start=1):
        flow.append(Paragraph(f"Questão {idx}", title_style))
        html = html_render_math_to_img(q.enunciado or "")
        soup = BeautifulSoup(html, 'lxml')
        accum_text = []
        for element in soup.recursiveChildGenerator():
            if getattr(element, 'name', None) == 'img' and element.get('src','').startswith('data:image/png;base64,'):
                if accum_text:
                    flow.append(Paragraph(''.join(accum_text).replace('\n','<br/>'), normal))
                    accum_text = []
                b64 = element['src'].split(',',1)[1]
                img_bytes = base64.b64decode(b64)
                flow.append(ReportLabImage(io.BytesIO(img_bytes), width=200, height=0, kind='proportional'))
            elif isinstance(element, str):
                text = element
                if text:
                    accum_text.append(text)
        if accum_text:
            flow.append(Paragraph(''.join(accum_text).replace('\n','<br/>'), normal))
        flow.append(Spacer(1, 10))
    return flow

@api_view(["POST"]) 
def generate_test_docx(request):
    ids = request.data.get('question_ids', [])
    if not isinstance(ids, list) or not ids:
        return Response({"detail": "question_ids deve ser uma lista não vazia"}, status=400)
    questions = list(Questao.objects.filter(id__in=ids))
    if not questions:
        return Response({"detail": "Nenhuma questão encontrada"}, status=404)
    document = Document()
    _docx_add_if_header(document, "PROVA", "Banco de Questões")
    for idx, q in enumerate(questions, start=1):
        document.add_paragraph(f"Questão {idx}")
        html = html_render_math_to_img(q.enunciado or "")
        soup = BeautifulSoup(html, 'lxml')
        accum_text = []
        for element in soup.recursiveChildGenerator():
            if getattr(element, 'name', None) == 'img' and element.get('src','').startswith('data:image/png;base64,'):
                if accum_text:
                    document.add_paragraph(''.join(accum_text))
                    accum_text = []
                b64 = element['src'].split(',',1)[1]
                img_bytes = base64.b64decode(b64)
                stream = io.BytesIO(img_bytes)
                document.add_picture(stream)
            elif isinstance(element, str):
                accum_text.append(element)
        if accum_text:
            document.add_paragraph(''.join(accum_text))
        if getattr(q, 'resposta', None):
            document.add_paragraph("Resposta:")
            html_r = html_render_math_to_img(q.resposta or "")
            soup_r = BeautifulSoup(html_r, 'lxml')
            accum_text_r = []
            for element in soup_r.recursiveChildGenerator():
                if getattr(element, 'name', None) == 'img' and element.get('src','').startswith('data:image/png;base64,'):
                    if accum_text_r:
                        document.add_paragraph(''.join(accum_text_r))
                        accum_text_r = []
                    b64 = element['src'].split(',',1)[1]
                    img_bytes = base64.b64decode(b64)
                    stream = io.BytesIO(img_bytes)
                    document.add_picture(stream)
                elif isinstance(element, str):
                    accum_text_r.append(element)
            if accum_text_r:
                document.add_paragraph(''.join(accum_text_r))
        document.add_paragraph("")
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    resp = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    resp["Content-Disposition"] = f'attachment; filename="prova_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx"'
    return resp

@api_view(["POST"]) 
def generate_test_pdf(request):
    ids = request.data.get('question_ids', [])
    if not isinstance(ids, list) or not ids:
        return Response({"detail": "question_ids deve ser uma lista não vazia"}, status=400)
    questions = list(Questao.objects.filter(id__in=ids))
    if not questions:
        return Response({"detail": "Nenhuma questão encontrada"}, status=404)
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=40, rightMargin=40, topMargin=160, bottomMargin=40)
    flow = _build_questions_flow(questions)
    styles = getSampleStyleSheet()
    normal = ParagraphStyle(name='Normal', parent=styles['Normal'], fontName='Helvetica', fontSize=11, leading=14, alignment=TA_JUSTIFY)
    title_style = ParagraphStyle(name='QTitle', parent=styles['Heading4'], spaceAfter=6)
    for idx, q in enumerate(questions, start=1):
        if getattr(q, 'resposta', None):
            flow.append(Paragraph(f"Resposta da Questão {idx}", title_style))
            html = html_render_math_to_img(q.resposta or "")
            soup = BeautifulSoup(html, 'lxml')
            accum_text = []
            for element in soup.recursiveChildGenerator():
                if getattr(element, 'name', None) == 'img' and element.get('src','').startswith('data:image/png;base64,'):
                    if accum_text:
                        flow.append(Paragraph(''.join(accum_text).replace('\n','<br/>'), normal))
                        accum_text = []
                    b64 = element['src'].split(',',1)[1]
                    img_bytes = base64.b64decode(b64)
                    flow.append(ReportLabImage(io.BytesIO(img_bytes), width=200, height=0, kind='proportional'))
                elif isinstance(element, str):
                    accum_text.append(element)
            if accum_text:
                flow.append(Paragraph(''.join(accum_text).replace('\n','<br/>'), normal))
            flow.append(Spacer(1, 10))
    def on_first_page(canvas, doc):
        _pdf_if_header(canvas, "PROVA", "Banco de Questões")
    doc.build(flow, onFirstPage=on_first_page, onLaterPages=on_first_page)
    pdf_value = buffer.getvalue()
    buffer.close()
    resp = HttpResponse(pdf_value, content_type='application/pdf')
    resp["Content-Disposition"] = f'attachment; filename="prova_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf"'
    return resp
@csrf_exempt
def upload_image(request):
    if request.method == 'POST' and request.FILES.get('upload'):
        image = request.FILES['upload']
        path = default_storage.save(f"uploads/{image.name}", image)
        url = default_storage.url(path)
        return JsonResponse({'url': url})
    return JsonResponse({'error': 'Invalid request'}, status=400)


@api_view(["POST"])
def signup(request):
    email = request.data.get("email")
    password = request.data.get("password")

    if not email or not password:
        return Response({"detail": "Campos obrigatórios"}, status=400)

    if User.objects.filter(username=email).exists():
        return Response({"detail": "Email já cadastrado"}, status=400)

    user = User.objects.create_user(username=email, email=email, password=password)
    return Response({"message": "Usuário criado com sucesso"}, status=201)


@api_view(["POST"])
def login_view(request):
    email = request.data.get("email")
    password = request.data.get("password")

    user = authenticate(username=email, password=password)
    if user is None:
        return Response({"detail": "Credenciais inválidas"}, status=400)

    refresh = RefreshToken.for_user(user)
    return Response({
        "token": str(refresh.access_token),
        "refresh": str(refresh),
        "user": {"email": user.email}
    })

def cadastro_questao(request):
    if request.method == 'POST':
        form = QuestaoForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            messages.success(request, "Questão cadastrada com sucesso!")
            return redirect('cadastro_questao')
        else:
            messages.error(request, "Erro ao cadastrar questão.")
    else:
        form = QuestaoForm()

    areas = Conteudo.objects.filter(tipo='area')
    return render(request, 'app/cadastro_questao.html', {'form': form, 'areas': areas})



def buscar_conteudos_filho(request):
    pai_id = request.GET.get('pai_id')
    if not pai_id:
        return JsonResponse({'error': 'pai_id ausente'}, status=400)

    filhos = Conteudo.objects.filter(pai_id=pai_id).values('id', 'nome')
    return JsonResponse(list(filhos), safe=False)

@api_view(["GET"])
def list_conteudos(request):
    """List all Conteudos, optionally filtered by tipo and pai_id"""
    conteudos = Conteudo.objects.all()
    
    tipo = request.query_params.get('tipo', None)
    if tipo:
        conteudos = conteudos.filter(tipo=tipo)
    
    pai_id = request.query_params.get('pai_id', None)
    if pai_id:
        conteudos = conteudos.filter(pai_id=pai_id)
    else:
        # If no pai_id, return only root-level (no parent) conteudos
        if tipo:
            conteudos = conteudos.filter(pai__isnull=True)
    
    serializer = ConteudoSerializer(conteudos, many=True)
    return Response(serializer.data)

@api_view(["GET"])
def get_unique_values(request):
    """Get unique values for fields like banca, tipo_questao, dificuldade, ano"""
    field = request.query_params.get('field', None)
    
    if not field:
        return Response({"error": "Field parameter is required"}, status=400)
    
    if field == 'banca':
        values = Questao.objects.values_list('banca', flat=True).distinct().order_by('banca')
    elif field == 'tipo_questao':
        values = Questao.objects.values_list('tipo_questao', flat=True).distinct().order_by('tipo_questao')
    elif field == 'dificuldade':
        values = Questao.objects.values_list('dificuldade', flat=True).distinct().order_by('dificuldade')
    elif field == 'ano':
        values = Questao.objects.values_list('ano', flat=True).distinct().order_by('-ano')
    else:
        return Response({"error": "Invalid field"}, status=400)
    
    return Response(list(values))

def inicio(request):
    return render(request, 'app/inicio.html')

@api_view(["GET"])
def list_questoes(request):
    questions = Questao.objects.all()
    serializer = QuestaoSerializer(questions, many=True)
    return Response(serializer.data)

@api_view(["GET"])
def questao_detail(request, pk):
    try:
        question = Questao.objects.get(pk=pk)
    except Questao.DoesNotExist:
        return Response({"detail": "Questão não encontrada"}, status=404)
    serializer = QuestaoSerializer(question)
    return Response(serializer.data)


class QuestaoList(generics.ListAPIView):
    queryset = Questao.objects.all()
    serializer_class = QuestaoSerializer

class QuestaoDetail(generics.RetrieveAPIView):
    queryset = Questao.objects.all()
    serializer_class = QuestaoSerializer

class QuestaoViewSet(viewsets.ModelViewSet):
    queryset = Questao.objects.all()
    serializer_class = QuestaoSerializer
    
    def get_queryset(self):
        queryset = Questao.objects.all()
        
        # Text search on enunciado
        search = self.request.query_params.get('search', None)
        if search:
            queryset = queryset.filter(enunciado__icontains=search)
        
        # Filter by area (by ID)
        area_id = self.request.query_params.get('area_id', None)
        if area_id:
            try:
                queryset = queryset.filter(area_id=int(area_id))
            except ValueError:
                pass
        
        # Filter by unidade (by ID)
        unidade_id = self.request.query_params.get('unidade_id', None)
        if unidade_id:
            try:
                queryset = queryset.filter(unidade_id=int(unidade_id))
            except ValueError:
                pass
        
        # Filter by topico (by ID)
        topico_id = self.request.query_params.get('topico_id', None)
        if topico_id:
            try:
                queryset = queryset.filter(topico_id=int(topico_id))
            except ValueError:
                pass
        
        # Filter by subtopico (by ID)
        subtopico_id = self.request.query_params.get('subtopico_id', None)
        if subtopico_id:
            try:
                queryset = queryset.filter(subtopico_id=int(subtopico_id))
            except ValueError:
                pass
        
        # Filter by categoria (by ID)
        categoria_id = self.request.query_params.get('categoria_id', None)
        if categoria_id:
            try:
                queryset = queryset.filter(categoria_id=int(categoria_id))
            except ValueError:
                pass
        
        # Filter by ano
        ano = self.request.query_params.get('ano', None)
        if ano:
            try:
                ano_int = int(ano)
                queryset = queryset.filter(ano=ano_int)
            except ValueError:
                pass
        
        # Filter by banca
        banca = self.request.query_params.get('banca', None)
        if banca:
            queryset = queryset.filter(banca__icontains=banca)
        
        # Filter by tipo_questao
        tipo_questao = self.request.query_params.get('tipo_questao', None)
        if tipo_questao:
            queryset = queryset.filter(tipo_questao__icontains=tipo_questao)
        
        # Filter by dificuldade
        dificuldade = self.request.query_params.get('dificuldade', None)
        if dificuldade:
            queryset = queryset.filter(dificuldade__icontains=dificuldade)
        
        return queryset