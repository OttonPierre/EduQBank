import io
from bs4 import BeautifulSoup
from django.http import HttpResponse
from rest_framework.decorators import api_view
from rest_framework.response import Response
from docx import Document
from datetime import datetime
from app.models import Questao
from app.utils import html_render_math_to_img
from .helpers import (
    _generate_with_pypandoc,
    _generate_with_pandoc,
    _docx_add_if_header,
    _get_image_bytes_from_html,
)


@api_view(["POST"]) 
def print_test_docx(request):
    """
    Gera um arquivo DOCX de prova a partir de uma lista de questões.

    Parâmetros aceitos em request.data:
      - question_ids: lista de IDs de questões (obrigatório)
      - gabarito_option (opcional, string):
            "apos_cada_questao"
            "final_arquivo"
            "somente_questoes"
            "somente_gabarito"
            "somente_gabarito_com_expectativa"
      - include_gabarito / use_resposta_gabarito:
            mantidos por compatibilidade com versões antigas do frontend.
    """
    ids = request.data.get('question_ids', [])

    # Compatibilidade: valores antigos ainda podem ser enviados
    include_gabarito = bool(request.data.get('include_gabarito', True))
    use_resposta_gabarito = bool(request.data.get('use_resposta_gabarito', False))
    gabarito_option = request.data.get('gabarito_option')

    # Se a opção de gabarito foi explicitamente informada pelo frontend novo,
    # convertemos para os flags internos.
    if gabarito_option:
        if gabarito_option == "somente_questoes":
            include_gabarito = False
            use_resposta_gabarito = False
        elif gabarito_option == "somente_gabarito":
            include_gabarito = True
            use_resposta_gabarito = True  # usar resposta_gabarito (letras, etc.)
        elif gabarito_option == "somente_gabarito_com_expectativa":
            include_gabarito = True
            use_resposta_gabarito = False  # usar expectativa / resposta completa
        elif gabarito_option in ("apos_cada_questao", "final_arquivo"):
            # Provas completas com gabarito em posições diferentes.
            include_gabarito = True
            # Para essas opções seguimos o valor enviado ou o padrão (False).
        else:
            # Valor desconhecido: assumir comportamento padrão atual
            gabarito_option = "final_arquivo"
    else:
        # Frontend antigo: inferir opção a partir dos flags
        if not include_gabarito:
            gabarito_option = "somente_questoes"
        else:
            gabarito_option = "final_arquivo"
    
    if not isinstance(ids, list) or not ids:
        return Response({"detail": "question_ids deve ser uma lista não vazia"}, status=400)
    questions = list(Questao.objects.filter(id__in=ids))
    if not questions:
        return Response({"detail": "Nenhuma questão encontrada"}, status=404)
    
    # Try pypandoc first (preferred)
    pp_bytes = _generate_with_pypandoc(
        questions,
        'docx',
        include_gabarito=include_gabarito,
        use_resposta_gabarito=use_resposta_gabarito,
        gabarito_option=gabarito_option,
    )
    if pp_bytes:
        resp = HttpResponse(pp_bytes, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        filename = "prova" if not include_gabarito else "prova_com_gabarito"
        resp["Content-Disposition"] = f'attachment; filename="{filename}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx"'
        return resp
    
    # Fallback to pandoc CLI
    pandoc_bytes = _generate_with_pandoc(
        questions,
        'docx',
        include_gabarito=include_gabarito,
        use_resposta_gabarito=use_resposta_gabarito,
        gabarito_option=gabarito_option,
    )
    if pandoc_bytes:
        resp = HttpResponse(pandoc_bytes, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        filename = "prova" if not include_gabarito else "prova_com_gabarito"
        resp["Content-Disposition"] = f'attachment; filename="{filename}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx"'
        return resp

    # Fallback to python-docx
    document = Document()
    _docx_add_if_header(document, "PROVA", "Banco de Questões")
    
    # Primeira página: apenas enunciados
    for idx, q in enumerate(questions, start=1):
        document.add_paragraph(f"Questão {idx}")
        html = html_render_math_to_img(q.enunciado or "")
        soup = BeautifulSoup(html, 'lxml')
        accum_text = []
        for element in soup.recursiveChildGenerator():
            if getattr(element, 'name', None) == 'img':
                src = element.get('src','')
                if accum_text:
                    document.add_paragraph(''.join(accum_text))
                    accum_text = []
                img_bytes = _get_image_bytes_from_html(src)
                if img_bytes:
                    stream = io.BytesIO(img_bytes)
                    document.add_picture(stream)
            elif isinstance(element, str):
                accum_text.append(element)
        if accum_text:
            document.add_paragraph(''.join(accum_text))
        document.add_paragraph("")
    
    # Segunda página: gabarito (se solicitado)
    if include_gabarito:
        document.add_page_break()
        _docx_add_if_header(document, "GABARITO", "Banco de Questões")
        for idx, q in enumerate(questions, start=1):
            document.add_paragraph(f"Questão {idx}")
            
            # Escolher entre resposta_gabarito ou resposta
            if use_resposta_gabarito and getattr(q, 'resposta_gabarito', None):
                # Usar resposta_gabarito (HTML rico)
                html_g = html_render_math_to_img(q.resposta_gabarito or "")
                soup_g = BeautifulSoup(html_g, 'lxml')
                accum_text_g = []
                for element in soup_g.recursiveChildGenerator():
                    if getattr(element, 'name', None) == 'img':
                        src = element.get('src','')
                        if accum_text_g:
                            document.add_paragraph(''.join(accum_text_g))
                            accum_text_g = []
                        img_bytes = _get_image_bytes_from_html(src)
                        if img_bytes:
                            stream = io.BytesIO(img_bytes)
                            document.add_picture(stream)
                    elif isinstance(element, str):
                        accum_text_g.append(element)
                if accum_text_g:
                    document.add_paragraph(''.join(accum_text_g))
            elif getattr(q, 'resposta', None):
                # Usar resposta (HTML rico)
                html_r = html_render_math_to_img(q.resposta or "")
                soup_r = BeautifulSoup(html_r, 'lxml')
                accum_text_r = []
                for element in soup_r.recursiveChildGenerator():
                    if getattr(element, 'name', None) == 'img':
                        src = element.get('src','')
                        if accum_text_r:
                            document.add_paragraph(''.join(accum_text_r))
                            accum_text_r = []
                        img_bytes = _get_image_bytes_from_html(src)
                        if img_bytes:
                            stream = io.BytesIO(img_bytes)
                            document.add_picture(stream)
                    elif isinstance(element, str):
                        accum_text_r.append(element)
                if accum_text_r:
                    document.add_paragraph(''.join(accum_text_r))
            document.add_paragraph("")
    
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    filename = "prova" if not include_gabarito else "prova_com_gabarito"
    resp = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    resp["Content-Disposition"] = f'attachment; filename="{filename}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx"'
    return resp

