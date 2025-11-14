import os
import base64
import tempfile
import subprocess
from typing import Optional
from bs4 import BeautifulSoup
from django.conf import settings
from urllib.parse import urlparse
from urllib.request import urlopen
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    import pypandoc  # type: ignore
except Exception:
    pypandoc = None  # type: ignore


def _convert_ckeditor_math_to_latex(html: str) -> str:
    r"""
    Converte fórmulas do CKEditor 4 (MathJax) para formato LaTeX que o Pandoc entende.
    
    O CKEditor salva fórmulas como:
    - <span class="math-tex">\(...\)</span> ou <span class="math-tex">\[...\]</span>
    - <script type="math/tex">...</script> ou <script type="math/tex; mode=display">...</script>
    
    O Pandoc precisa de:
    - \(...\) para fórmulas inline
    - \[...\] para fórmulas display
    
    Retorna apenas o conteúdo HTML interno, sem tags <html><body>.
    """
    if not html:
        return html
    
    soup = BeautifulSoup(html, 'lxml')
    changed = False
    
    # Converter <span class="math-tex"> para LaTeX delimiters
    for span in soup.find_all('span', class_=lambda c: c and 'math-tex' in (c if isinstance(c, str) else ' '.join(c)).split()):
        latex_content = span.get_text(strip=True)
        if latex_content:
            # Remover delimiters existentes se houver
            clean_latex = latex_content.strip()
            if clean_latex.startswith('\\(') and clean_latex.endswith('\\)'):
                clean_latex = clean_latex[2:-2].strip()
            elif clean_latex.startswith('\\[') and clean_latex.endswith('\\]'):
                clean_latex = clean_latex[2:-2].strip()
            elif clean_latex.startswith('$') and clean_latex.endswith('$'):
                clean_latex = clean_latex[1:-1].strip()
            elif clean_latex.startswith('$$') and clean_latex.endswith('$$'):
                clean_latex = clean_latex[2:-2].strip()
            
            # Determinar se é display ou inline baseado no contexto
            # Se o span está em um parágrafo ou dentro de texto, é inline
            parent = span.parent
            is_display = parent and parent.name in ['div', 'p'] and span == parent.contents[0] and len(parent.contents) == 1
            
            if is_display:
                latex_str = soup.new_string(f'\\[{clean_latex}\\]')
            else:
                latex_str = soup.new_string(f'\\({clean_latex}\\)')
            
            span.replace_with(latex_str)
            changed = True
    
    # Converter <script type="math/tex"> para LaTeX delimiters
    for script in soup.find_all('script'):
        script_type = (script.get('type') or '').lower()
        if script_type.startswith('math/tex'):
            latex_content = script.get_text(strip=True)
            if latex_content:
                # Remover delimiters existentes se houver
                clean_latex = latex_content.strip()
                if clean_latex.startswith('\\(') and clean_latex.endswith('\\)'):
                    clean_latex = clean_latex[2:-2].strip()
                elif clean_latex.startswith('\\[') and clean_latex.endswith('\\]'):
                    clean_latex = clean_latex[2:-2].strip()
                elif clean_latex.startswith('$') and clean_latex.endswith('$'):
                    clean_latex = clean_latex[1:-1].strip()
                elif clean_latex.startswith('$$') and clean_latex.endswith('$$'):
                    clean_latex = clean_latex[2:-2].strip()
                
                # Determinar se é display ou inline
                is_display = 'mode=display' in script_type or 'display' in script_type
                
                if is_display:
                    latex_str = soup.new_string(f'\\[{clean_latex}\\]')
                else:
                    latex_str = soup.new_string(f'\\({clean_latex}\\)')
                
                script.replace_with(latex_str)
                changed = True
    
    if changed:
        # Retornar apenas o conteúdo interno (sem <html><body>)
        if soup.body:
            return soup.body.decode_contents()
        elif soup.contents:
            return ''.join(str(c) for c in soup.contents)
        else:
            return str(soup)
    
    # Se não houve mudanças, retornar o conteúdo interno mesmo assim
    if soup.body:
        return soup.body.decode_contents()
    return html


def _rewrite_img_src_to_fs_paths(html: str) -> str:
    """Rewrite <img src> that point to MEDIA_URL into absolute filesystem paths so Pandoc can embed them."""
    soup = BeautifulSoup(html, 'lxml')
    media_url = getattr(settings, 'MEDIA_URL', '/media/')
    media_root = getattr(settings, 'MEDIA_ROOT', None)
    for img in soup.find_all('img'):
        src = img.get('src', '')
        if not src:
            continue
        parsed = urlparse(src)
        if src.startswith('data:image'):
            continue
        if not parsed.scheme and media_root and src.startswith(media_url):
            rel_path = src[len(media_url):]
            abs_path = os.path.join(media_root, rel_path)
            if os.path.exists(abs_path):
                img['src'] = abs_path
    return str(soup)


def _generate_with_pypandoc(questions, out_format: str, include_gabarito: bool = True, use_resposta_gabarito: bool = False) -> Optional[bytes]:
    """Use pypandoc to convert HTML (with LaTeX) to PDF/DOCX using installed Pandoc and MiKTeX. Returns bytes or None."""
    if pypandoc is None:
        return None
    try:
        head = (
            "<meta charset='utf-8'/>"
            "<style>body{font-family: Arial, sans-serif;margin:40px;} img{max-width:100%;}</style>"
        )
        # Primeira página: apenas enunciados
        body_parts = [
            "<h2 style='text-align:center'>INSTITUTO FEDERAL – Sistema de Avaliação</h2>",
            "<h1 style='text-align:center'>PROVA</h1>",
            "<div>Professor(a): __________________ &nbsp;&nbsp; Turma: ______ &nbsp;&nbsp; Data: ____/____/______ &nbsp;&nbsp; Nota: ______</div>",
            "<div>Aluno(a): _________________________________________________________________</div>",
        ]
        for idx, q in enumerate(questions, start=1):
            body_parts.append(f"<h3>Questão {idx}</h3>")
            # Converter fórmulas do CKEditor para LaTeX que Pandoc entende
            enunciado = _convert_ckeditor_math_to_latex(q.enunciado or '')
            body_parts.append(f"<div>{enunciado}</div>")

        if include_gabarito:
            body_parts.append("<h1 style='text-align:center; page-break-before: always;'>GABARITO</h1>")
            
            for idx, q in enumerate(questions, start=1):
                body_parts.append(f"<h3>Questão {idx}</h3>")
                # Escolher entre resposta_gabarito ou resposta
                if use_resposta_gabarito and getattr(q, 'resposta_gabarito', None):
                    resposta_g = _convert_ckeditor_math_to_latex(q.resposta_gabarito or '')
                    body_parts.append(f"<div>{resposta_g}</div>")
                elif getattr(q, 'resposta', None):
                    resposta = _convert_ckeditor_math_to_latex(q.resposta or '')
                    body_parts.append(f"<div>{resposta}</div>")
        html = f"<html><head>{head}</head><body>{''.join(body_parts)}</body></html>"
        html = _rewrite_img_src_to_fs_paths(html)

        reader = 'html+tex_math_dollars+tex_math_single_backslash'
        extra_args = []
        media_root = getattr(settings, 'MEDIA_ROOT', None)
        if media_root:
            extra_args.extend(['--resource-path', media_root])

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = os.path.join(tmpdir, f'output.{out_format}')
            pypandoc.convert_text(html, to=out_format, format=reader, outputfile=output_path, extra_args=extra_args)
            with open(output_path, 'rb') as f:
                return f.read()
    except Exception:
        return None


def _generate_with_pandoc(questions, out_format: str, include_gabarito: bool = True, use_resposta_gabarito: bool = False) -> Optional[bytes]:
    """Try to use pandoc (and MiKTeX/LaTeX engine) to produce PDF/DOCX with real formulas. Returns bytes or None on failure."""
    try:
        head = (
            "<meta charset='utf-8'/>"
            "<style>body{font-family: Arial, sans-serif;margin:40px;} img{max-width:100%;}</style>"
        )
        # Primeira página: apenas enunciados
        body_parts = [
            "<h2 style='text-align:center'>INSTITUTO FEDERAL – Sistema de Avaliação</h2>",
            "<h1 style='text-align:center'>PROVA</h1>",
            "<div>Professor(a): __________________ &nbsp;&nbsp; Turma: ______ &nbsp;&nbsp; Data: ____/____/______ &nbsp;&nbsp; Nota: ______</div>",
            "<div>Aluno(a): _________________________________________________________________</div>",
        ]
        for idx, q in enumerate(questions, start=1):
            body_parts.append(f"<h3>Questão {idx}</h3>")
            # Converter fórmulas do CKEditor para LaTeX que Pandoc entende
            enunciado = _convert_ckeditor_math_to_latex(q.enunciado or '')
            body_parts.append(f"<div>{enunciado}</div>")
        
        # Segunda página: gabarito (se solicitado)
        if include_gabarito:
            body_parts.append("<p style='page-break-before: always;'></p>")
            body_parts.append("<h2 style='text-align:center'>INSTITUTO FEDERAL – Sistema de Avaliação</h2>")
            body_parts.append("<h1 style='text-align:center'>GABARITO</h1>")
            body_parts.append("<div>Professor(a): __________________ &nbsp;&nbsp; Turma: ______ &nbsp;&nbsp; Data: ____/____/______ &nbsp;&nbsp; Nota: ______</div>")
            body_parts.append("<div>Aluno(a): _________________________________________________________________</div>")
            
            for idx, q in enumerate(questions, start=1):
                body_parts.append(f"<h3>Questão {idx}</h3>")
                # Escolher entre resposta_gabarito ou resposta
                if use_resposta_gabarito and getattr(q, 'resposta_gabarito', None):
                    resposta_g = _convert_ckeditor_math_to_latex(q.resposta_gabarito or '')
                    body_parts.append(f"<div>{resposta_g}</div>")
                elif getattr(q, 'resposta', None):
                    resposta = _convert_ckeditor_math_to_latex(q.resposta or '')
                    body_parts.append(f"<div>{resposta}</div>")
        html = f"<html><head>{head}</head><body>{''.join(body_parts)}</body></html>"
        html = _rewrite_img_src_to_fs_paths(html)

        with tempfile.TemporaryDirectory() as tmpdir:
            in_path = os.path.join(tmpdir, 'input.html')
            out_path = os.path.join(tmpdir, f'output.{out_format}')
            with open(in_path, 'w', encoding='utf-8') as f:
                f.write(html)
            
            reader = 'html+tex_math_dollars+tex_math_single_backslash'
            cmd = ['pandoc', in_path, '-f', reader, '-o', out_path]
            media_root = getattr(settings, 'MEDIA_ROOT', None)
            if media_root:
                cmd.extend(['--resource-path', media_root])
            
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            with open(out_path, 'rb') as f:
                return f.read()
    except Exception:
        return None


def _docx_add_if_header(document: Document, title: str = "PROVA", subtitle: str = "") -> None:
    section = document.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    run = header_para.add_run("INSTITUTO FEDERAL – Sistema de Avaliação\n")
    run.font.size = Pt(10)
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph()
    r = p.add_run(title)
    r.font.size = Pt(16)
    r.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        p2 = document.add_paragraph(subtitle)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info = document.add_paragraph("Professor(a): __________________    Turma: ______    Data: ____/____/______    Nota: ______")
    info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph("Aluno(a): _________________________________________________________________")
    document.add_paragraph("")


def _get_image_bytes_from_html(src: str) -> Optional[bytes]:
    """Extract image bytes from HTML src attribute (supports data URIs, local media, and URLs)."""
    if src.startswith('data:image') and 'base64,' in src:
        b64 = src.split(',',1)[1]
        return base64.b64decode(b64)
    try:
        parsed = urlparse(src)
        if not parsed.scheme:
            media_url = getattr(settings, 'MEDIA_URL', '/media/')
            media_root = getattr(settings, 'MEDIA_ROOT', None)
            if media_root and src.startswith(media_url):
                rel_path = src[len(media_url):]
                abs_path = os.path.join(media_root, rel_path)
                with open(abs_path, 'rb') as fh:
                    return fh.read()
        else:
            with urlopen(src) as resp:
                return resp.read()
    except Exception:
        pass
    return None

